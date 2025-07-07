import os
import sys
import json
from docx import Document
from docx2pdf import convert
from datetime import datetime
import google.generativeai as genai

# ==== Nhập đầu vào ====
if len(sys.argv) != 5:
    print("❌ Cách dùng: python pdf.py <profileName> <ten_ngan_hang> <don_vi> <nhan_vien>")
    sys.exit(1)

profile_name = sys.argv[1]
ten_ngan_hang = sys.argv[2]
don_vi = sys.argv[3]
nhan_vien = sys.argv[4]

# ==== Cấu hình Gemini ====
genai.configure(api_key="##############")
model = genai.GenerativeModel("gemini-2.5-flash-preview-05-20")

# ==== Load dữ liệu JSON ====
json_path = os.path.join("final_analysis", profile_name, "final_output.json")
if not os.path.exists(json_path):
    print("❌ Không tìm thấy file JSON:", json_path)
    sys.exit(1)

with open(json_path, "r", encoding="utf-8") as f:
    extracted_data = json.load(f)

# ==== Chuẩn bị nội dung prompt ====
prompt = f"""
Dựa trên dữ liệu sau, hãy tạo một **Tờ trình thẩm định tín dụng** theo mẫu chuẩn của ngân hàng, được thực hiện bởi nhân viên {nhan_vien} thuộc chi nhánh {don_vi} của ngân hàng {ten_ngan_hang}. Hãy hành văn trang trọng, chuẩn mẫu ngân hàng, đầy đủ bố cục, định dạng như một tài liệu Word chuyên nghiệp.

Dữ liệu đầu vào (JSON):
{json.dumps(extracted_data, indent=2, ensure_ascii=False)}

Các phần nội dung cần trình bày như sau:

---

### 1. THÔNG TIN KHÁCH HÀNG
- Tên doanh nghiệp, mã số thuế
- Địa chỉ trụ sở chính
- Đại diện pháp luật, chức vụ
- Ngành nghề kinh doanh
- Thời gian hoạt động
- Nhu cầu vay vốn và mục đích sử dụng vốn
- Loại hình khách hàng (mới/hiện hữu)
- Phân loại ngành, nhóm khách hàng

---

### 2. HIỆN TRẠNG & KIẾN NGHỊ CẤP TÍN DỤNG
- Tổng hợp hạn mức vay đang có và kiến nghị hạn mức mới
- Thông tin về tài sản bảo đảm:
  - Loại tài sản, địa chỉ, tình trạng pháp lý
  - Giá trị định giá, tỷ lệ cho vay
  - Quan hệ tài sản với khách hàng
- Thẩm quyền phê duyệt: đề xuất các bên tham gia xét duyệt

---

### 3. TÌNH HÌNH QUAN HỆ TÍN DỤNG VỚI CÁC TCTD & ĐÁNH GIÁ CIC
- Đánh giá lịch sử tín dụng tại ngân hàng hiện tại và các TCTD khác
- Dư nợ tại các TCTD, xếp hạng tín dụng, khả năng thanh toán
- Đánh giá CIC cập nhật theo ngày kiểm tra gần nhất
- Tình trạng tín dụng và CIC của các cá nhân là cổ đông/lãnh đạo liên quan

---

### 4. PHÂN TÍCH HOẠT ĐỘNG KINH DOANH & TÀI CHÍNH
- Tổng hợp các chỉ tiêu tài chính theo từng năm:
  - Doanh thu, lợi nhuận, tổng tài sản, nợ phải trả
  - ROA, ROE, tỷ lệ vay/NHĐT/NHTS
- Nhận xét về tình hình kinh doanh:
  - Tăng trưởng doanh thu, năng lực cạnh tranh, quy mô thị trường
  - Lợi nhuận ròng, cơ cấu tài chính, khả năng thanh toán ngắn hạn/dài hạn
- So sánh tỷ lệ nợ/vốn CSH, khả năng trả nợ, xu hướng tài chính

---

### 5. PHƯƠNG ÁN SỬ DỤNG VỐN & QUAN HỆ VỚI NGÂN HÀNG
- Mô tả mục đích sử dụng vốn: khoản vay cụ thể, kế hoạch kinh doanh
- Pháp lý khoản vay: phù hợp với ngành nghề kinh doanh
- Pháp lý đầu tư: pháp lý các tài sản đầu tư hoặc các dự án
- Kết quả chấm điểm tín dụng nội bộ
- Đánh giá tiềm năng phát triển và mục tiêu duy trì quan hệ
- Sản phẩm vay, kỳ hạn, hạn mức, thu nhập dự kiến từ quan hệ tín dụng

---

### 6. TÓM TẮT RỦI RO & BIỆN PHÁP KIỂM SOÁT
- Danh sách các rủi ro tiềm ẩn:
  - Rủi ro mất khả năng thanh toán
  - Rủi ro về tài sản bảo đảm
  - Rủi ro pháp lý
  - Rủi ro cạnh tranh,...
- Biện pháp kiểm soát rủi ro tương ứng do ngân hàng đề xuất

---

### 7. ĐỀ XUẤT TÍN DỤNG & KẾT LUẬN
- Nhận xét tổng thể của cán bộ thẩm định
- Đề xuất cấp tín dụng:
  - Tổng mức cấp tín dụng đề nghị
  - Sản phẩm vay, thời hạn vay, phương thức trả gốc/lãi
  - Tài sản đảm bảo (tên, giá trị, tỷ lệ cho vay,...)
  - Điều kiện cấp tín dụng: điều kiện giải ngân, thanh toán nợ, sử dụng vốn
- Ghi chú các ràng buộc: bảo hiểm TSĐB, chuyển doanh thu, thời hạn giải ngân,...
- Đề xuất các bộ phận liên quan kiểm soát tín dụng: P.L.3, P.KSNB, TSĐB, Scoring,...

---

### 8. CHỮ KÝ & PHÊ DUYỆT
- Chữ ký các bên:
  - Nhân viên thực hiện
  - Trưởng bộ phận/HCB
  - Giám đốc chi nhánh
- Ý kiến của bộ phận kiểm soát nội bộ
- Bảng tích chọn đơn vị phụ trách phê duyệt (UBTD, BTD, Scoring,...)

---

✅ Hãy tạo một tài liệu hoàn chỉnh dựa trên dữ liệu JSON ở trên. Hành văn theo phong cách chuyên nghiệp, súc tích, chuẩn văn bản ngân hàng. Có thể sử dụng bảng biểu (table) nếu cần để thể hiện thông tin rõ ràng.
"""


response = model.generate_content(prompt)
generated_text = response.text.strip()

# ==== Tạo file Word ====
doc = Document()
doc.add_heading('TỜ TRÌNH THẨM ĐỊNH TÍN DỤNG', level=0)

doc.add_paragraph(f'Ngân hàng: {ten_ngan_hang}')
doc.add_paragraph(f'Đơn vị: {don_vi}')
doc.add_paragraph(f'Người thực hiện: {nhan_vien}')
doc.add_paragraph(f'Ngày: {datetime.now().strftime("%d/%m/%Y")}')

doc.add_paragraph("")  # Dòng trống
for line in generated_text.split("\n"):
    if line.strip():
        doc.add_paragraph(line.strip())

# ==== Lưu file Word và chuyển sang PDF ====
output_folder = os.path.join("final_analysis", profile_name)
word_path = os.path.join(output_folder, "to_trinh.docx")
pdf_path = os.path.join(output_folder, "to_trinh.pdf")

doc.save(word_path)

# Chuyển sang PDF
try:
    convert(word_path, pdf_path)
    print("", pdf_path)
except Exception as e:
    print("Error:", e)
    sys.exit(1)
